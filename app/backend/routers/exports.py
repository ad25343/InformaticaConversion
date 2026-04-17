# Copyright (c) 2026 ad25343 — https://github.com/ad25343/InformaticaConversion
# Licensed under CC BY-NC 4.0. Commercial use requires written permission.
"""
Exports sub-router: file download endpoints for converted output.
"""
from __future__ import annotations

from fastapi import APIRouter, File, UploadFile, HTTPException
from fastapi.responses import StreamingResponse, JSONResponse
from typing import Optional as _Opt

from ._helpers import (
    db, logger,
    _validate_job_id,
    s2t_excel_path,
    validate_upload_size,
    build_output_zip,
    JobStatus,
)

router = APIRouter(prefix="")


# ─────────────────────────────────────────────
# Internal helpers
# ─────────────────────────────────────────────

def _safe_filename(name: str) -> str:
    """Sanitise a mapping name for use in a download filename."""
    return "".join(c if c.isalnum() or c in "-_" else "_" for c in name)


def _load_manifest_overrides_from_bytes(xlsx_bytes: bytes) -> list:
    """Write bytes to a temp file, parse overrides, clean up, return list."""
    from ..agents import manifest_agent
    import tempfile as _tempfile
    import os as _os

    with _tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
        tmp.write(xlsx_bytes)
        tmp_path = tmp.name
    try:
        return manifest_agent.load_overrides(tmp_path)
    finally:
        _os.unlink(tmp_path)


def _validate_manifest_job_for_upload(job, job_id: str) -> None:
    """Raise HTTPException if the job is missing or not in awaiting_review state."""
    if not job:
        _validate_job_id(job_id)
        raise HTTPException(404, "Job not found")
    if job["status"] != JobStatus.AWAITING_REVIEW.value:
        raise HTTPException(
            400,
            f"Manifest overrides can only be uploaded while the job is awaiting review "
            f"(current status: {job['status']}). Download the manifest, annotate it, "
            f"then upload it before submitting your sign-off.",
        )


# ─────────────────────────────────────────────
# Download converted code
# ─────────────────────────────────────────────

@router.get("/jobs/{job_id}/s2t/download")
async def download_s2t_excel(job_id: str):
    """Download the Source-to-Target mapping Excel workbook for a job."""
    job = await db.get_job(job_id)
    if not job:
        _validate_job_id(job_id)
        raise HTTPException(404, "Job not found")

    path = s2t_excel_path(job_id)
    if not path or not path.exists():
        raise HTTPException(404, "S2T Excel file not found — the job may not have completed Step 2 yet")

    from fastapi.responses import FileResponse
    return FileResponse(
        str(path),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename=path.name,
    )


@router.get("/jobs/{job_id}/manifest.xlsx")
async def download_manifest_xlsx(job_id: str):
    """
    Generate and return the pre-conversion mapping manifest xlsx on demand.
    The manifest is NOT stored in state (too large); it is regenerated from the
    graph dict each time this endpoint is called.
    """
    job = await db.get_job(job_id)
    if not job:
        _validate_job_id(job_id)
        raise HTTPException(404, "Job not found")
    graph = job.get("state", {}).get("graph")
    if not graph:
        raise HTTPException(404, "Manifest not available — job has not completed parsing yet")

    from ..agents import manifest_agent
    import io as _io
    report = manifest_agent.build_manifest(graph)
    xlsx_bytes = manifest_agent.write_xlsx_bytes(report)
    safe = job.get("filename", "mapping").replace(".xml", "").replace(" ", "_")
    return StreamingResponse(
        _io.BytesIO(xlsx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="manifest_{safe}.xlsx"'},
    )


@router.post("/jobs/{job_id}/manifest-upload")
async def upload_manifest_overrides(job_id: str, file: UploadFile = File(...)):
    """
    Accept a reviewer-annotated manifest xlsx and store the overrides in job state.

    The reviewer downloads the manifest via GET /jobs/{job_id}/manifest.xlsx,
    fills in the 'Reviewer Override' column on the 'Review Required' sheet for any
    LOW or UNMAPPED rows, then re-uploads the annotated file here.

    Must be called while the job is at Gate 1 (awaiting_review).
    The conversion agent picks up the stored overrides when the pipeline resumes
    after sign-off, resolving lineage gaps before generating code.
    """
    job = await db.get_job(job_id)
    _validate_manifest_job_for_upload(job, job_id)

    fname = (file.filename or "").lower()
    if not fname.endswith(".xlsx"):
        raise HTTPException(400, "Manifest file must be a .xlsx file")

    xlsx_bytes = await file.read()
    validate_upload_size(xlsx_bytes, label=file.filename)
    if not xlsx_bytes:
        raise HTTPException(400, "Uploaded manifest file is empty")

    overrides = _load_manifest_overrides_from_bytes(xlsx_bytes)
    overrides_dicts = [o.model_dump() for o in overrides]

    await db.update_job(
        job_id, JobStatus.AWAITING_REVIEW.value, 5,
        {"manifest_overrides": overrides_dicts},
    )

    logger.info("Manifest overrides uploaded: job_id=%s override_count=%d",
                job_id, len(overrides_dicts))

    return {
        "message": f"Manifest uploaded successfully. {len(overrides_dicts)} override(s) stored.",
        "job_id": job_id,
        "override_count": len(overrides_dicts),
        "overrides": overrides_dicts,
    }


@router.get("/jobs/{job_id}/download/{filename}")
async def download_file(job_id: str, filename: str):
    job = await db.get_job(job_id)
    if not job:
        _validate_job_id(job_id)
        raise HTTPException(404, "Job not found")
    conversion = job["state"].get("conversion", {})
    files = conversion.get("files", {})
    if filename not in files:
        raise HTTPException(404, f"File '{filename}' not found in conversion output")

    # GAP #14 — Validate the filename is safe before serving
    import pathlib
    _safe_name = pathlib.PurePosixPath(filename).name
    _ALLOWED_EXTS = {".py", ".sql", ".yaml", ".yml", ".txt", ".md", ".json", ".sh", ".cfg", ".ini", ".toml"}
    _ext = pathlib.PurePosixPath(_safe_name).suffix.lower()
    if _ext not in _ALLOWED_EXTS:
        logger.warning("Blocked download of disallowed extension: job=%s filename=%s", job_id, filename)
        raise HTTPException(400, f"File extension '{_ext}' is not permitted for download.")

    return JSONResponse({"filename": filename, "content": files[filename]})


@router.get("/jobs/{job_id}/tests/download/{filename:path}")
async def download_test_file(job_id: str, filename: str):
    """Download a generated test file by path (e.g. tests/test_conversion.py)."""
    job = await db.get_job(job_id)
    if not job:
        _validate_job_id(job_id)
        raise HTTPException(404, "Job not found")
    test_report = job["state"].get("test_report", {})
    files = test_report.get("test_files", {})
    if filename not in files:
        raise HTTPException(404, f"Test file '{filename}' not found")

    import pathlib
    _ALLOWED_EXTS = {".py", ".sql", ".yaml", ".yml", ".txt", ".md", ".json", ".sh", ".cfg", ".ini", ".toml"}
    _ext = pathlib.PurePosixPath(filename).suffix.lower()
    if _ext not in _ALLOWED_EXTS:
        logger.warning("Blocked test download of disallowed extension: job=%s filename=%s", job_id, filename)
        raise HTTPException(400, f"File extension '{_ext}' is not permitted for download.")

    return JSONResponse({"filename": filename, "content": files[filename]})


# ─────────────────────────────────────────────
# Output ZIP Download (v2.5.0)
# ─────────────────────────────────────────────

@router.get("/jobs/{job_id}/output.zip")
async def download_output_zip(job_id: str):
    """
    Download all generated conversion output files as a ZIP archive.

    Bundles every file from state["conversion"]["files"] into a single
    ZIP preserving folder structure.  Built directly from DB state so it
    works regardless of whether the job folder has been written to disk.

    Only available for jobs that have reached AWAITING_CODE_REVIEW or
    COMPLETE status (i.e. conversion has run).
    """
    job = await db.get_job(job_id)
    if not job:
        _validate_job_id(job_id)
        raise HTTPException(404, "Job not found")

    state = job.get("state", {})
    conversion = state.get("conversion", {})
    files = conversion.get("files", {})
    if not files:
        raise HTTPException(404, "No output files found — conversion has not completed for this job.")

    zip_bytes = build_output_zip(state)
    mapping_name = conversion.get("mapping_name", job_id)
    filename = f"{_safe_filename(mapping_name)}_output.zip"

    return StreamingResponse(
        iter([zip_bytes]),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─────────────────────────────────────────────
# DOCX export for analyst docs (3a / 3b)
# ─────────────────────────────────────────────

def _md_to_docx_bytes(md_text: str, title: str, mapping_name: str = "", tier: str = "") -> bytes:
    """Convert markdown text to a polished, visually appealing DOCX document."""
    import io as _io
    import re as _re
    from datetime import datetime
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    # Colour palette
    C_NAVY    = RGBColor(0x1e, 0x3a, 0x5f)
    C_DARK    = RGBColor(0x1e, 0x29, 0x3b)
    C_BODY    = RGBColor(0x33, 0x41, 0x55)
    C_MUTED   = RGBColor(0x64, 0x74, 0x8b)
    C_CODE    = RGBColor(0x4f, 0x46, 0xe5)   # indigo
    C_WARN    = RGBColor(0x92, 0x40, 0x0e)   # amber-800
    C_WHITE   = RGBColor(0xff, 0xff, 0xff)

    HEADER_BG = "1e3a5f"
    STRIPE_BG = "f8fafc"
    BORDER_CLR = "cbd5e1"

    doc = Document()

    # ── Global styles ──
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(10)
    style_normal.font.color.rgb = C_BODY
    style_normal.paragraph_format.space_after = Pt(5)
    style_normal.paragraph_format.line_spacing = 1.25

    for lvl, (sz, clr) in {0: (22, C_DARK), 1: (16, C_NAVY), 2: (13, C_NAVY), 3: (11, C_BODY)}.items():
        hs = doc.styles[f"Heading {lvl + 1}" if lvl else "Title"]
        hs.font.name = "Calibri"
        hs.font.size = Pt(sz)
        hs.font.color.rgb = clr
        hs.font.bold = True

    # ── Margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Cover / Title block ──
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(4)
    run = p_title.add_run(title)
    run.font.size = Pt(22)
    run.font.color.rgb = C_DARK
    run.bold = True

    if mapping_name:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(2)
        run = p_sub.add_run(mapping_name)
        run.font.size = Pt(11)
        run.font.color.rgb = C_MUTED

    # Metadata line
    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(8)
    p_meta.paragraph_format.space_after = Pt(16)
    meta_parts = [f"Generated: {datetime.now().strftime('%B %d, %Y')}"]
    if tier:
        meta_parts.append(f"Complexity: {tier.upper()}")
    meta_parts.append("Informatica Conversion Platform")
    run = p_meta.add_run("  |  ".join(meta_parts))
    run.font.size = Pt(8)
    run.font.color.rgb = C_MUTED

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(12)
    pPr = p_div._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="{HEADER_BG}"/></w:pBdr>')
    pPr.append(pBdr)

    # ── Helper: add inline-formatted runs to a paragraph ──
    def _add_rich_runs(para, text):
        parts = _re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                r = para.add_run(part[2:-2])
                r.bold = True
                r.font.color.rgb = C_DARK
            elif part.startswith("`") and part.endswith("`"):
                r = para.add_run(part[1:-1])
                r.font.name = "Consolas"
                r.font.size = Pt(9)
                r.font.color.rgb = C_CODE
            else:
                r = para.add_run(part)

    # ── Helper: style a table ──
    def _style_table(table, rows):
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        # Set border on table
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:color="{BORDER_CLR}"/>'
            f'<w:left w:val="single" w:sz="4" w:color="{BORDER_CLR}"/>'
            f'<w:bottom w:val="single" w:sz="4" w:color="{BORDER_CLR}"/>'
            f'<w:right w:val="single" w:sz="4" w:color="{BORDER_CLR}"/>'
            f'<w:insideH w:val="single" w:sz="4" w:color="{BORDER_CLR}"/>'
            f'<w:insideV w:val="single" w:sz="4" w:color="{BORDER_CLR}"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

        ncols = max(len(r) for r in rows)
        for ri, row_cells in enumerate(rows):
            for ci, cell_text in enumerate(row_cells):
                if ci >= ncols:
                    continue
                cell = table.rows[ri].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

                # Header row
                if ri == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEADER_BG}" w:val="clear"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                    r = p.add_run(cell_text.upper() if len(cell_text) < 30 else cell_text)
                    r.bold = True
                    r.font.size = Pt(7.5)
                    r.font.color.rgb = C_WHITE
                    r.font.name = "Calibri"
                # Striped body rows
                elif ri % 2 == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{STRIPE_BG}" w:val="clear"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                    _add_rich_runs(p, cell_text)
                    for r in p.runs:
                        r.font.size = Pt(8.5)
                else:
                    _add_rich_runs(p, cell_text)
                    for r in p.runs:
                        r.font.size = Pt(8.5)

    # ── Main document body ──
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Headings
        if stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=3)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=2)
            h.paragraph_format.space_before = Pt(20)
            h.paragraph_format.space_after = Pt(8)
            # Add subtle bottom border
            pPr = h._p.get_or_add_pPr()
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="2" w:color="{BORDER_CLR}"/></w:pBdr>')
            pPr.append(pBdr)
        elif stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=1)
            h.paragraph_format.space_before = Pt(24)
        elif stripped.startswith("---"):
            pass  # skip

        # Tables
        elif stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1

            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip("|").split("|")]
                if all(_re.match(r"^[-:]+$", c) for c in cells):
                    continue
                rows.append(cells)

            if rows:
                ncols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncols)
                _style_table(table, rows)

        # Code blocks
        elif stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1

            if lang == "mermaid":
                # Build node ID → label map from all lines first
                node_labels = {}
                for cl in code_lines:
                    for nm in _re.finditer(r'(\w+)\[([^\]]+)\]', cl):
                        node_labels[nm.group(1)] = nm.group(2)

                def _lbl(nid):
                    return node_labels.get(nid, nid)

                # Parse flow connections
                flow_steps = []
                for cl in code_lines:
                    cl = cl.strip()
                    if not cl or cl.startswith("graph ") or cl.startswith("%%"):
                        continue
                    # Normalize arrow syntax to extract: src, edge_label, tgt
                    # Handles: -->, -.-> , -.text.->, ==>, -->|label|
                    edge_label = None
                    # Extract dotted-arrow text like -.lookup.->
                    dot_m = _re.search(r'-\.([^.]+)\.->', cl)
                    if dot_m:
                        edge_label = dot_m.group(1).strip()
                    # Normalize all arrow variants to -->
                    normalized = _re.sub(r'-+(?:\.[^.]*\.)?->?', '-->', cl)
                    normalized = _re.sub(r'==+>', '-->', normalized)
                    arrow_m = _re.match(
                        r'\s*(\w+)(?:\[[^\]]*\])?\s*-->\s*(?:\|([^|]*)\|)?\s*(\w+)(?:\[[^\]]*\])?',
                        normalized
                    )
                    if arrow_m:
                        src = _lbl(arrow_m.group(1))
                        edge_label = edge_label or arrow_m.group(2)
                        tgt = _lbl(arrow_m.group(3))
                        if edge_label:
                            flow_steps.append(f"{src}  \u2192  {tgt}  ({edge_label})")
                        else:
                            flow_steps.append(f"{src}  \u2192  {tgt}")

                # Render as styled flow box
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                pPr = p._p.get_or_add_pPr()
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'<w:top w:val="single" w:sz="6" w:color="{HEADER_BG}"/>'
                    f'<w:left w:val="single" w:sz="6" w:color="{HEADER_BG}"/>'
                    f'<w:bottom w:val="single" w:sz="6" w:color="{HEADER_BG}"/>'
                    f'<w:right w:val="single" w:sz="6" w:color="{HEADER_BG}"/>'
                    f'</w:pBdr>'
                )
                pPr.append(pBdr)
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="f0f4ff" w:val="clear"/>')
                pPr.append(shd)
                r = p.add_run("  \u25B6  Data Flow Diagram")
                r.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = C_NAVY

                if flow_steps:
                    for step in flow_steps:
                        sp = doc.add_paragraph(style="List Bullet")
                        sp.paragraph_format.space_after = Pt(1)
                        sr = sp.add_run(step)
                        sr.font.size = Pt(8.5)
                        sr.font.color.rgb = C_BODY
                else:
                    # Fallback: show mermaid source formatted
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.space_after = Pt(6)
                    r2 = p2.add_run("\n".join(code_lines))
                    r2.font.name = "Consolas"
                    r2.font.size = Pt(7.5)
                    r2.font.color.rgb = C_MUTED
            else:
                # Styled code block with left accent bar
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                pPr = p._p.get_or_add_pPr()
                pBdr = parse_xml(
                    f'<w:pBdr {nsdecls("w")}>'
                    f'<w:left w:val="single" w:sz="16" w:space="8" w:color="6366f1"/>'
                    f'<w:top w:val="single" w:sz="4" w:color="{BORDER_CLR}"/>'
                    f'<w:bottom w:val="single" w:sz="4" w:color="{BORDER_CLR}"/>'
                    f'<w:right w:val="single" w:sz="4" w:color="{BORDER_CLR}"/>'
                    f'</w:pBdr>'
                )
                pPr.append(pBdr)
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="f8fafc" w:val="clear"/>')
                pPr.append(shd)
                if lang and lang != "mermaid":
                    r_lang = p.add_run(f"  {lang.upper()}\n")
                    r_lang.font.size = Pt(6.5)
                    r_lang.font.color.rgb = C_MUTED
                    r_lang.bold = True
                r = p.add_run("\n".join(code_lines))
                r.font.name = "Consolas"
                r.font.size = Pt(8)
                r.font.color.rgb = C_BODY

        # Blockquotes (warnings / notes)
        elif stripped.startswith(">"):
            text = _re.sub(r"^>\s*", "", stripped)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:left w:val="single" w:sz="16" w:space="8" w:color="f59e0b"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="fffbeb" w:val="clear"/>')
            pPr.append(shd)
            _add_rich_runs(p, text)
            for r in p.runs:
                r.font.size = Pt(9)

        # Bullet lists
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_runs(p, text)

        # Numbered lists
        elif _re.match(r"^\d+\.\s", stripped):
            text = _re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _add_rich_runs(p, text)

        # Regular paragraph with inline formatting
        elif stripped:
            p = doc.add_paragraph()
            _add_rich_runs(p, stripped)

        i += 1

    # ── Footer ──
    p_foot = doc.add_paragraph()
    p_foot.paragraph_format.space_before = Pt(24)
    pPr = p_foot._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="6" w:space="4" w:color="{BORDER_CLR}"/></w:pBdr>')
    pPr.append(pBdr)
    r = p_foot.add_run(f"Generated by Informatica Conversion Tool  |  {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  Confidential")
    r.font.size = Pt(7.5)
    r.font.color.rgb = C_MUTED

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/jobs/{job_id}/doc/{doc_key}.docx")
async def download_doc_docx(job_id: str, doc_key: str):
    """Download analyst_summary_md, analyst_view_md, or analyst_gaps_md as a styled DOCX."""
    import io as _io

    ALLOWED = {
        "analyst_summary": ("analyst_summary_md", "Analyst Summary",        "analyst_summary"),
        "analyst_view":    ("analyst_view_md",     "Technical Specification","technical_specification"),
        "analyst_gaps":    ("analyst_gaps_md",     "Gaps & Review Findings", "gaps_review"),
    }
    if doc_key not in ALLOWED:
        raise HTTPException(400, f"Invalid doc_key: {doc_key}. Allowed: {list(ALLOWED)}")

    state_key, title, suffix = ALLOWED[doc_key]
    job = await db.get_job(job_id)
    if not job:
        _validate_job_id(job_id)
        raise HTTPException(404, "Job not found")

    state = job.get("state", {})
    md_text = state.get(state_key, "")
    if not md_text:
        raise HTTPException(404, f"{state_key} not available for this job.")

    safe = _safe_filename(job.get("filename", job_id).replace(".xml", ""))
    mapping_name = job.get("filename", "").replace(".xml", "")
    tier = state.get("complexity", {}).get("tier", "")
    docx_bytes = _md_to_docx_bytes(md_text, title, mapping_name=mapping_name, tier=tier)

    return StreamingResponse(
        _io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{safe}_{suffix}.docx"'},
    )
